"""DocxBlockParser: reads .docx via python-docx and produces ParsedDocument."""
from __future__ import annotations

import hashlib
import re
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph
from tabulate import tabulate as _tabulate

from src.models import (
    Block,
    BlockLocation,
    BlockType,
    ExtractionMethod,
    FileType,
    ParseError,
    ParsedDocument,
    TableRecord,
)

# Compiled patterns
_RE_REVISION_TRIGGER = re.compile(r"^\d+\.\d+\.\d+版修訂項目如下")
_RE_DATE_END = re.compile(r"^\d+\.\d+\.\d+.*(訂定|修訂)$")
_RE_CHAPTER = re.compile(r"^第[一二三四五六七八九十百]+章[\s\u3000]+")
_RE_SUBSECTION = re.compile(r"^第[\s\u3000]*[一二三四五六七八九十百]+[\s\u3000]*節")
_RE_ARTICLE = re.compile(r"^第[\s\u3000]*([一二三四五六七八九十百]+)[\s\u3000]*條[\s\u3000]+")
_RE_POINT = re.compile(r"^([一二三四五六七八九十]+)、")
# Require actual parentheses — avoids matching "一、xxx" sub-point style
_RE_CATEGORY = re.compile(r"^\([一二三四五六七八九十]\)[\s\u3000]*\S")
_RE_TOC_ENTRY = re.compile(r"^.+(\t|[ \u3000]{2,})\d{1,4}\s*$")
_RE_HEADING_STYLE = re.compile(r"Heading\s*(\d+)", re.IGNORECASE)
_REGULATION_ENDS = re.compile(r"(辦法|要點|規範|準則|程序|計畫|規則|注意事項|規定|要項|政策|措施|方法|指引)$")
_REGULATION_STARTS = re.compile(r"^(汐止區農會|農會)")


def _iter_container(container, doc: Document):
    """Yield w:p and w:tbl children of a container element (body or sdtContent)."""
    for child in container.iterchildren():
        if child.tag == qn("w:p"):
            yield Paragraph(child, doc)
        elif child.tag == qn("w:tbl"):
            yield Table(child, doc)
        elif child.tag == qn("w:sdt"):
            # SDT (Structured Document Tag) — e.g., Word auto-generated TOC
            sdt_content = child.find(qn("w:sdtContent"))
            if sdt_content is not None:
                yield from _iter_container(sdt_content, doc)


def iterate_body_in_order(doc: Document):
    """Yield paragraphs and tables in original document order, including SDT content."""
    yield from _iter_container(doc.element.body, doc)


def _is_bold(para: Paragraph) -> bool:
    return any(run.bold for run in para.runs if run.text.strip())


def _is_content_start(text: str, is_bold: bool) -> bool:
    """Return True if a bold paragraph is clearly regulation content, not revision history."""
    if not is_bold:
        return False
    if _RE_CATEGORY.match(text) and len(text) < 30:
        return True
    if _REGULATION_STARTS.match(text) and _REGULATION_ENDS.search(text):
        return True
    if _RE_CHAPTER.match(text):
        return True
    return False


def _is_italic(para: Paragraph) -> bool:
    return any(run.italic for run in para.runs if run.text.strip())


def _classify_paragraph(para: Paragraph, text: str, is_bold: bool) -> BlockType:
    # Rule 4: REGULATION_TITLE
    if is_bold and _REGULATION_STARTS.match(text) and _REGULATION_ENDS.search(text):
        return BlockType.REGULATION_TITLE

    # Rule 5: ESTABLISHMENT_DATE / REVISION_DATE
    m = _RE_DATE_END.match(text)
    if m:
        return BlockType.ESTABLISHMENT_DATE if m.group(1) == "訂定" else BlockType.REVISION_DATE

    # Rule 7: TOC_ENTRY — must come before structural headings so that TOC lines
    # like「第一章 通則\t1」are not misclassified as CHAPTER_HEADING.
    # DATE and REGULATION_TITLE are safe above this because they never carry tab+pagenum.
    if _RE_TOC_ENTRY.match(text):
        return BlockType.TOC_ENTRY

    # Rule 8: CATEGORY_HEADING (moved after TOC_ENTRY)
    if (
        is_bold
        and len(text) < 30
        and _RE_CATEGORY.match(text)
        and not _RE_TOC_ENTRY.match(text)
    ):
        return BlockType.CATEGORY_HEADING

    # Rule 9: CHAPTER_HEADING
    if _RE_CHAPTER.match(text):
        return BlockType.CHAPTER_HEADING

    # Rule 10: SUBSECTION_HEADING
    if _RE_SUBSECTION.match(text):
        return BlockType.SUBSECTION_HEADING

    # Rule 11: ARTICLE
    if _RE_ARTICLE.match(text):
        return BlockType.ARTICLE

    # Rule 12: POINT
    if _RE_POINT.match(text):
        return BlockType.POINT

    # Rule 12: Heading style — but regulation title pattern takes priority
    style_name = ""
    try:
        style_name = para.style.name or ""
    except Exception:
        pass
    if _RE_HEADING_STYLE.search(style_name):
        # Heading styles suppress run-level bold; check regulation pattern explicitly
        if _REGULATION_STARTS.match(text) and _REGULATION_ENDS.search(text):
            return BlockType.REGULATION_TITLE
        return BlockType.SECTION_HEADING

    return BlockType.PARAGRAPH


def _extract_article_number(text: str) -> str | None:
    m = _RE_ARTICLE.match(text)
    return m.group(1) if m else None


def _extract_point_number(text: str) -> str | None:
    m = _RE_POINT.match(text)
    return m.group(1) if m else None


def _extract_heading_level(para: Paragraph) -> int:
    try:
        style_name = para.style.name or ""
        m = _RE_HEADING_STYLE.search(style_name)
        if m:
            return int(m.group(1))
    except Exception:
        pass
    return 1


def _try_merge_wrapped_title(current_text: str, next_text: str | None) -> str | None:
    """If current_text is the first line of a split regulation title, merge with next_text.

    Handles cases where Word wraps a long title across two consecutive paragraphs.
    Returns the merged title, or None if the merge condition is not met.
    """
    if not next_text:
        return None
    if not _REGULATION_STARTS.match(current_text):
        return None
    # First line must NOT already end with a legal suffix (would already be complete)
    if _REGULATION_ENDS.search(current_text):
        return None
    # Second line must be a short completion (not a paragraph of body text)
    if len(next_text) > 30:
        return None
    if not _REGULATION_ENDS.search(next_text):
        return None
    merged = current_text.rstrip() + next_text.lstrip()
    if not (10 <= len(merged) <= 60):
        return None
    return merged


def _table_to_raw_grid(table: Table) -> tuple[list[list[str]], int, int]:
    grid: list[list[str]] = []
    for row in table.rows:
        row_data: list[str] = []
        seen_tc: set[int] = set()
        for cell in row.cells:
            tc_id = id(cell._tc)
            if tc_id in seen_tc:
                continue
            seen_tc.add(tc_id)
            cell_text = "\n".join(p.text for p in cell.paragraphs)
            row_data.append(cell_text)
        grid.append(row_data)

    n_rows = len(grid)
    n_cols = max((len(r) for r in grid), default=0)
    for row in grid:
        while len(row) < n_cols:
            row.append("")
    return grid, n_rows, n_cols


def _table_to_markdown(raw_grid: list[list[str]]) -> str:
    if not raw_grid:
        return ""
    if len(raw_grid) == 1:
        return _tabulate([], headers=raw_grid[0], tablefmt="github")
    return _tabulate(raw_grid[1:], headers=raw_grid[0], tablefmt="github")


def _table_to_html(table: Table) -> str:
    parts = ["<table>"]
    for row_idx, row in enumerate(table.rows):
        parts.append("<tr>")
        seen_tc: set[int] = set()
        for cell in row.cells:
            tc_id = id(cell._tc)
            if tc_id in seen_tc:
                continue
            seen_tc.add(tc_id)
            inner = "<br>".join(p.text for p in cell.paragraphs if p.text.strip())
            tag = "th" if row_idx == 0 else "td"
            parts.append(f"<{tag}>{inner}</{tag}>")
        parts.append("</tr>")
    parts.append("</table>")
    return "\n".join(parts)


def _detect_merged_cells(table: Table) -> list[tuple[int, int, int, int]]:
    # TODO: implement proper merged cell detection
    return []


def _fill_context_snippets(blocks: list[Block], max_lookaround: int = 3) -> None:
    content_types = {
        BlockType.PARAGRAPH,
        BlockType.CHAPTER_HEADING,
        BlockType.SUBSECTION_HEADING,
        BlockType.ARTICLE,
        BlockType.POINT,
        BlockType.REGULATION_TITLE,
        BlockType.CATEGORY_HEADING,
        BlockType.LIST_ITEM,
    }

    for i, block in enumerate(blocks):
        for j in range(i - 1, max(i - 1 - max_lookaround, -1), -1):
            if blocks[j].block_type in content_types and blocks[j].content.strip():
                block.location.preceding_text_snippet = blocks[j].content[:80]
                break

        for j in range(i + 1, min(i + 1 + max_lookaround, len(blocks))):
            if blocks[j].block_type in content_types and blocks[j].content.strip():
                block.location.following_text_snippet = blocks[j].content[:80]
                break


def parse_docx(file_path: Path) -> ParsedDocument:
    """Read .docx and produce ParsedDocument.

    Responsibilities:
    - Read paragraphs + tables in original order to produce blocks
    - Fill complete location on every block (except structural path fields)
    - Produce raw_grid / markdown / html for every table
    - TABLE blocks in the content flow are pointer-only (table_ref)

    Does NOT handle:
    - Document shape detection
    - Virtual document splitting
    - Structural path fields (filled by CompilationSplitter)
    """
    file_path = Path(file_path)

    with open(file_path, "rb") as f:
        file_bytes = f.read()
    source_hash = hashlib.sha256(file_bytes).hexdigest()

    doc = Document(file_path)

    blocks: list[Block] = []
    table_records: list[TableRecord] = []
    parse_errors: list[ParseError] = []

    doc_order_idx = 0
    para_idx = 0
    table_idx = 0

    # Revision block accumulation state
    revision_lines: list[str] = []
    revision_doc_idx: int = 0
    revision_para_idx: int = 0

    def flush_revision() -> None:
        nonlocal doc_order_idx
        if not revision_lines:
            return
        content = "\n".join(revision_lines)
        location = BlockLocation(
            document_order_index=revision_doc_idx,
            source_paragraph_index=revision_para_idx,
        )
        blocks.append(
            Block(
                block_type=BlockType.REVISION_HISTORY_BLOCK,
                content=content,
                location=location,
            )
        )
        revision_lines.clear()

    elements = list(iterate_body_in_order(doc))
    n_elements = len(elements)
    i = 0

    while i < n_elements:
        element = elements[i]

        if isinstance(element, Paragraph):
            text = element.text
            stripped = text.strip()
            if not stripped:
                para_idx += 1
                i += 1
                continue

            bold = _is_bold(element)

            # Revision trigger — flush previous block and start new accumulation
            if _RE_REVISION_TRIGGER.match(stripped):
                flush_revision()
                revision_lines.append(stripped)
                revision_doc_idx = doc_order_idx
                revision_para_idx = para_idx
                doc_order_idx += 1
                para_idx += 1
                i += 1
                continue

            # 目錄 — flush revision state, emit TOC_HEADING
            if stripped == "目錄":
                flush_revision()
                loc = BlockLocation(
                    document_order_index=doc_order_idx,
                    source_paragraph_index=para_idx,
                )
                blocks.append(
                    Block(block_type=BlockType.TOC_HEADING, content=stripped, location=loc)
                )
                doc_order_idx += 1
                para_idx += 1
                i += 1
                continue

            # Still accumulating revision lines — but stop if we hit content start
            if revision_lines:
                if _is_content_start(stripped, bold):
                    flush_revision()
                    # Fall through to normal classification below
                else:
                    revision_lines.append(stripped)
                    para_idx += 1
                    i += 1
                    continue

            # Cross-line regulation title: try to merge with the immediately next paragraph
            next_para_text: str | None = None
            if i + 1 < n_elements and isinstance(elements[i + 1], Paragraph):
                _nxt = elements[i + 1].text.strip()
                if _nxt:
                    next_para_text = _nxt

            merged_title = _try_merge_wrapped_title(stripped, next_para_text)
            if merged_title:
                loc = BlockLocation(
                    document_order_index=doc_order_idx,
                    source_paragraph_index=para_idx,
                )
                blocks.append(
                    Block(
                        block_type=BlockType.REGULATION_TITLE,
                        content=merged_title,
                        location=loc,
                        is_bold=bold,
                        is_italic=_is_italic(element),
                    )
                )
                doc_order_idx += 1
                para_idx += 2  # consumed two source paragraphs
                i += 2  # skip the next element
                continue

            # Normal classification
            try:
                block_type = _classify_paragraph(element, stripped, bold)
                loc = BlockLocation(
                    document_order_index=doc_order_idx,
                    source_paragraph_index=para_idx,
                )

                article_number: str | None = None
                point_number: str | None = None
                heading_level: int | None = None

                if block_type == BlockType.ARTICLE:
                    article_number = _extract_article_number(stripped)
                elif block_type == BlockType.POINT:
                    point_number = _extract_point_number(stripped)
                elif block_type == BlockType.CHAPTER_HEADING:
                    heading_level = 2
                elif block_type == BlockType.SUBSECTION_HEADING:
                    heading_level = 3
                elif block_type == BlockType.SECTION_HEADING:
                    heading_level = _extract_heading_level(element)

                blocks.append(
                    Block(
                        block_type=block_type,
                        content=stripped,
                        location=loc,
                        is_bold=bold,
                        is_italic=_is_italic(element),
                        heading_level=heading_level,
                        article_number=article_number,
                        point_number=point_number,
                    )
                )
                doc_order_idx += 1
            except Exception as exc:
                parse_errors.append(
                    ParseError(
                        error_type="PARAGRAPH_PARSE_FAILED",
                        message=str(exc),
                        block_order_index=doc_order_idx,
                        severity="WARNING",
                    )
                )
                loc = BlockLocation(
                    document_order_index=doc_order_idx,
                    source_paragraph_index=para_idx,
                )
                blocks.append(
                    Block(block_type=BlockType.UNKNOWN, content=stripped, location=loc)
                )
                doc_order_idx += 1

            para_idx += 1

        elif isinstance(element, Table):
            # Tables always flush pending revision accumulation
            flush_revision()

            try:
                raw_grid, n_rows, n_cols = _table_to_raw_grid(element)
                md = _table_to_markdown(raw_grid)
                html = _table_to_html(element)
                merged = _detect_merged_cells(element)

                table_id = len(table_records)
                loc = BlockLocation(
                    document_order_index=doc_order_idx,
                    source_table_index=table_idx,
                )

                table_records.append(
                    TableRecord(
                        table_id=table_id,
                        raw_grid=raw_grid,
                        markdown=md,
                        html=html,
                        n_rows=n_rows,
                        n_cols=n_cols,
                        merged_cells=merged,
                        location=loc,
                        extraction_method=ExtractionMethod.PYTHON_DOCX_NATIVE,
                    )
                )

                content = f"[表格:{n_rows}列{n_cols}欄]"
                blocks.append(
                    Block(
                        block_type=BlockType.TABLE,
                        content=content,
                        location=loc,
                        table_ref=table_id,
                    )
                )
                doc_order_idx += 1
            except Exception as exc:
                parse_errors.append(
                    ParseError(
                        error_type="TABLE_PARSE_FAILED",
                        message=str(exc),
                        block_order_index=doc_order_idx,
                        severity="WARNING",
                    )
                )
                loc = BlockLocation(
                    document_order_index=doc_order_idx,
                    source_table_index=table_idx,
                )
                blocks.append(
                    Block(
                        block_type=BlockType.UNKNOWN,
                        content="[表格解析失敗]",
                        location=loc,
                    )
                )
                doc_order_idx += 1

            table_idx += 1

        i += 1

    # Flush any remaining revision block at end of document
    flush_revision()

    _fill_context_snippets(blocks)

    return ParsedDocument(
        source_path=str(file_path),
        source_hash=source_hash,
        file_type=FileType.DOCX,
        blocks=blocks,
        tables=table_records,
        parse_errors=parse_errors,
    )
